"""app/extraction/parsers/docx.py — DOCX parsing with python-docx"""
from __future__ import annotations

import io

import structlog

log = structlog.get_logger(__name__)


class DocxParser:
    """Parse DOCX bytes into structured text, sections, and tables."""

    def parse(self, file_bytes: bytes) -> dict:
        """
        Returns:
            {
                "raw_text": str,
                "sections": list[{"title": str, "text": str}],
                "tables": list[{"title": str, "columns": list, "rows": list[dict]}],
                "parser_used": "python-docx",
                "errors": list[str]
            }
        Never raises. Returns partial results with errors on failure.
        """
        try:
            from docx import Document

            doc = Document(io.BytesIO(file_bytes))

            raw_text = ""
            sections: list[dict] = []
            tables: list[dict] = []
            errors: list[str] = []

            current_section_title = "Introduction"
            current_section_text: list[str] = []

            for paragraph in doc.paragraphs:
                style_name = paragraph.style.name if paragraph.style else ""
                text = paragraph.text.strip()

                if style_name.startswith("Heading"):
                    # Save current section
                    if current_section_text:
                        sections.append(
                            {
                                "title": current_section_title,
                                "text": "\n".join(current_section_text).strip(),
                            }
                        )
                    current_section_title = text or current_section_title
                    current_section_text = []
                else:
                    if text:
                        raw_text += text + "\n"
                        current_section_text.append(text)

            # Save last section
            if current_section_text:
                sections.append(
                    {
                        "title": current_section_title,
                        "text": "\n".join(current_section_text).strip(),
                    }
                )

            # Extract tables
            for i, table in enumerate(doc.tables):
                try:
                    rows_raw = table.rows
                    if not rows_raw:
                        continue

                    # First row = headers
                    headers = [
                        cell.text.strip() or f"col_{j}"
                        for j, cell in enumerate(rows_raw[0].cells)
                    ]
                    data_rows = []
                    for row in rows_raw[1:]:
                        row_dict = {}
                        for j, cell in enumerate(row.cells):
                            key = headers[j] if j < len(headers) else f"col_{j}"
                            row_dict[key] = cell.text.strip()
                        if any(v for v in row_dict.values()):
                            data_rows.append(row_dict)

                    tables.append(
                        {
                            "title": f"Table {i + 1}",
                            "columns": headers,
                            "rows": data_rows,
                        }
                    )
                except Exception as te:
                    log.warning("docx_table_error", index=i, error=str(te))
                    errors.append(f"table_{i}_error: {te}")

            return {
                "raw_text": raw_text,
                "sections": sections,
                "tables": tables,
                "parser_used": "python-docx",
                "errors": errors,
            }

        except Exception as e:
            log.error("docx_parse_failed", error=str(e))
            return {
                "raw_text": "",
                "sections": [],
                "tables": [],
                "parser_used": "python-docx",
                "errors": [str(e)],
            }
